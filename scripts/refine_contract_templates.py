from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path("templates/contracts")
AGREEMENT = BASE / "sales_brokerage_agreement.docx"
RECEIPT = BASE / "payment_receipt.docx"


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def add_field(paragraph, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"
    run_text = OxmlElement("w:r")
    run_text.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = OxmlElement("w:r")
    run.append(begin)
    paragraph._element.append(run)

    run2 = OxmlElement("w:r")
    run2.append(instr)
    paragraph._element.append(run2)

    run3 = OxmlElement("w:r")
    run3.append(separate)
    paragraph._element.append(run3)

    paragraph._element.append(run_text)

    run4 = OxmlElement("w:r")
    run4.append(end)
    paragraph._element.append(run4)


def set_footer_page_fields(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph(p)
    p.add_run("Page ")
    add_field(p, " PAGE ")
    p.add_run(" of ")
    add_field(p, " NUMPAGES ")


def update_header(doc, reference_label_fr, reference_label_ar):
    header = doc.sections[0].header
    if not header.tables:
        return
    table = header.tables[0]

    center_cell = table.cell(0, 1)
    if center_cell.paragraphs:
        p_logo = center_cell.paragraphs[0]
        clear_paragraph(p_logo)
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run("[LOGO AXIRA / شعار أكسيرا]")
        run.bold = True

    right_cell = table.cell(0, 2)
    if len(right_cell.paragraphs) >= 2:
        p_ref = right_cell.paragraphs[0]
        ref_text = p_ref.text
        value = ref_text.split(":", 1)[1].strip() if ":" in ref_text else ""
        clear_paragraph(p_ref)
        p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ref.add_run(f"{reference_label_fr} / {reference_label_ar}: {value}").bold = True

        p_date = right_cell.paragraphs[1]
        date_text = p_date.text
        date_value = date_text.split(":", 1)[1].strip() if ":" in date_text else "{{contract_date}}"
        clear_paragraph(p_date)
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.add_run(f"Date / التاريخ: {date_value}").bold = True


def replace_text_in_paragraphs(doc, old, new):
    for p in doc.paragraphs:
        if old in p.text:
            full = p.text.replace(old, new)
            clear_paragraph(p)
            p.add_run(full)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        full = p.text.replace(old, new)
                        clear_paragraph(p)
                        p.add_run(full)


def update_agreement_title(doc):
    if not doc.paragraphs:
        return
    p = doc.paragraphs[0]
    text = p.text.strip()
    if "CONTRAT DE VENTE ET DE COURTAGE" in text and "عقد بيع و وساطة" not in text:
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run("CONTRAT DE VENTE ET DE COURTAGE")
        r1.bold = True
        p.add_run("\n")
        r2 = p.add_run("عقد بيع و وساطة")
        r2.bold = True


def refine_agreement():
    doc = Document(str(AGREEMENT))
    update_header(doc, "Référence du contrat", "مرجع العقد")
    set_footer_page_fields(doc)
    update_agreement_title(doc)
    replace_text_in_paragraphs(doc, "Entre les soussignes", "Entre les soussignés")
    doc.save(str(AGREEMENT))


def refine_receipt():
    doc = Document(str(RECEIPT))
    update_header(doc, "Numéro du reçu", "رقم الإيصال")
    set_footer_page_fields(doc)
    replace_text_in_paragraphs(doc, "Entre les soussignes", "Entre les soussignés")
    doc.save(str(RECEIPT))


if __name__ == "__main__":
    refine_agreement()
    refine_receipt()
