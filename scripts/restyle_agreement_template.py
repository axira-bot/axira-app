import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/axira/Desktop/axira-app")
TEMPLATE = ROOT / "templates/contracts/sales_brokerage_agreement.docx"
LOGO = ROOT / "public/axira-logo.png"

CHARCOAL = "1E1E1E"
GOLD = "C8A84B"

PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")
OPEN_SECTION_RE = re.compile(r"\{\{#[^}]+\}\}")
CLOSE_SECTION_RE = re.compile(r"\{\{\/[^}]+\}\}")
ARABIC_RE = re.compile(r"[\u0600-\u06FF]")


def read_docx_xml_text(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path, "r") as zf:
        parts = [n for n in zf.namelist() if n.startswith("word/") and n.endswith(".xml")]
        chunks = [zf.read(p).decode("utf-8", errors="ignore") for p in parts]
    return "\n".join(chunks)


def token_counts(docx_path: Path) -> dict:
    xml = read_docx_xml_text(docx_path)
    return {
        "placeholders_total": len(PLACEHOLDER_RE.findall(xml)),
        "open_sections": len(OPEN_SECTION_RE.findall(xml)),
        "close_sections": len(CLOSE_SECTION_RE.findall(xml)),
    }


def set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color_hex: str, size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def clear_header_footer(container):
    for p in list(container.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(container.tables):
        t._element.getparent().remove(t._element)


def style_run(run, bold_placeholders_only=False):
    text = run.text or ""
    is_ar = bool(ARABIC_RE.search(text))
    if is_ar:
        run.font.name = "Traditional Arabic"
        run._r.rPr.rFonts.set(qn("w:cs"), "Traditional Arabic")
        run.font.size = Pt(12)
    else:
        run.font.name = "Helvetica"
        run._r.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        run.font.size = Pt(11)
    if "{{" in text and "}}" in text:
        run.bold = True
    elif not bold_placeholders_only:
        run.bold = run.bold


def style_body(doc: Document):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            style_run(run, bold_placeholders_only=True)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        style_run(run, bold_placeholders_only=True)


def add_header(section):
    header = section.header
    clear_header_footer(header)

    header_table = header.add_table(rows=1, cols=2, width=Inches(7.1))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(3.2)
    header_table.columns[1].width = Inches(3.9)

    left = header_table.cell(0, 0)
    right = header_table.cell(0, 1)

    logo_p = left.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_p.add_run()
    logo_run.add_picture(str(LOGO), width=Inches(1.55))

    title_p = right.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = title_p.add_run("AXIRA TRADING FZE")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(30, 30, 30)
    style_run(r1)

    meta_p = right.add_paragraph(
        "Réf. contrat / مرجع العقد: {{contract_reference}}\nDate / التاريخ: {{contract_date}}"
    )
    meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in meta_p.runs:
        style_run(run, bold_placeholders_only=True)

    line_p = header.add_paragraph("")
    set_paragraph_bottom_border(line_p, GOLD, "8")

    accent_p = header.add_paragraph("")
    set_paragraph_bottom_border(accent_p, CHARCOAL, "4")


def add_footer(section):
    footer = section.footer
    clear_header_footer(footer)

    gold = footer.add_paragraph("")
    set_paragraph_bottom_border(gold, GOLD, "8")

    table = footer.add_table(rows=1, cols=2, width=Inches(7.1))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(3.55)
    table.columns[1].width = Inches(3.55)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    set_cell_shading(left, CHARCOAL)
    set_cell_shading(right, CHARCOAL)

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l1 = lp.add_run("AXIRA TRADING FZE\n")
    l1.bold = True
    l1.font.color.rgb = RGBColor(255, 255, 255)
    style_run(l1)
    l2 = lp.add_run("{{fze_phone}} | {{fze_email}}\n{{fze_address}}")
    l2.font.color.rgb = RGBColor(230, 230, 230)
    style_run(l2, bold_placeholders_only=True)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = rp.add_run("AXIRA AUTO (ALGÉRIE)\n")
    r1.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)
    style_run(r1)
    r2 = rp.add_run("{{auto_phone}} | {{auto_email}}\n{{auto_address}}")
    r2.font.color.rgb = RGBColor(230, 230, 230)
    style_run(r2, bold_placeholders_only=True)

    page_p = footer.add_paragraph("Page ")
    page_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_p.runs[0].font.size = Pt(9)
    fld_page = OxmlElement("w:fldSimple")
    fld_page.set(qn("w:instr"), "PAGE")
    page_p._p.append(fld_page)
    page_p.add_run(" of ")
    fld_num = OxmlElement("w:fldSimple")
    fld_num.set(qn("w:instr"), "NUMPAGES")
    page_p._p.append(fld_num)


def main():
    before = token_counts(TEMPLATE)
    doc = Document(str(TEMPLATE))

    normal = doc.styles["Normal"]
    normal.font.name = "Helvetica"
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    style_body(doc)

    for section in doc.sections:
        add_header(section)
        add_footer(section)

    doc.save(str(TEMPLATE))
    after = token_counts(TEMPLATE)

    with zipfile.ZipFile(TEMPLATE, "r") as zf:
        media_files = [n for n in zf.namelist() if n.startswith("word/media/")]
        has_logo_media = any(n.lower().endswith(".png") for n in media_files)
        header_rels = [
            zf.read(n).decode("utf-8", errors="ignore")
            for n in zf.namelist()
            if n.startswith("word/_rels/header") and n.endswith(".xml.rels")
        ]
        embedded = any('TargetMode="External"' not in rel for rel in header_rels)

    print(
        {
            "before": before,
            "after": after,
            "has_logo_media": has_logo_media,
            "header_relationship_embedded": embedded,
        }
    )


if __name__ == "__main__":
    main()
